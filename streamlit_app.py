import streamlit as st
import google.generativeai as genai
import os
import io
import json
import re
import numpy as np
import matplotlib.pyplot as plt
from io import StringIO
from pypdf import PdfReader
from docx import Document
from datetime import date, datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

# ==========================================
# 0. ADMINISTRATION & UTILISATEURS
# ==========================================
# Pour ajouter un utilisateur :
# 1. Génère un hash avec : import hashlib; hashlib.sha256("motdepasse".encode()).hexdigest()
# 2. Ajoute une ligne dans UTILISATEURS ci-dessous
#
# Rôles disponibles : "admin" (accès complet) | "user" (accès standard)

import hashlib

def hash_mdp(mdp):
    return hashlib.sha256(mdp.encode()).hexdigest()

UTILISATEURS = {
    "admin": {
        "nom": "Administrateur",
        "hash": hash_mdp("wisc-admin-2025"),   # ← Change ce mot de passe !
        "role": "admin"
    },
    "psychologue1": {
        "nom": "Dr. Dupont",
        "hash": hash_mdp("motdepasse1"),        # ← Change ce mot de passe !
        "role": "user"
    },
}

DATE_EXPIRATION = date(2026, 12, 31)


# ==========================================
# 1. CONFIGURATION & SÉCURITÉ
# ==========================================
st.set_page_config(
    page_title="Assistant WISC-V",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# 2. THÈME CSS PROFESSIONNEL
# ==========================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Libre+Baskerville:ital,wght@0,400;0,700;1,400&family=DM+Sans:opsz,wght@9..40,300;9..40,400;9..40,500;9..40,600&display=swap');

:root {
    --bleu-nuit:    #0D1B2A;
    --bleu-marine:  #1B3A5C;
    --bleu-moyen:   #2D6A9F;
    --or-medical:   #C9A84C;
    --or-clair:     #E8C97A;
    --blanc-casse:  #F5F2EE;
    --gris-doux:    #E8E4DF;
    --gris-texte:   #6B7280;
    --rouge-alerte: #B94040;
    --vert-ok:      #2E7D5A;
    --fond-app:     #F0EDE8;
}

html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif !important;
    color: var(--bleu-nuit) !important;
}

.stApp {
    background-color: var(--fond-app) !important;
    background-image:
        radial-gradient(ellipse at 10% 0%, rgba(29,58,92,0.06) 0%, transparent 60%),
        radial-gradient(ellipse at 90% 100%, rgba(201,168,76,0.05) 0%, transparent 60%);
}

/* ---- Titres ---- */
h1 {
    font-family: 'Libre Baskerville', serif !important;
    font-size: 1.85rem !important;
    color: var(--bleu-marine) !important;
    letter-spacing: -0.02em !important;
    font-weight: 700 !important;
    border-bottom: 2px solid var(--or-medical);
    padding-bottom: 0.4rem;
    margin-bottom: 1rem !important;
}
h2 {
    font-family: 'Libre Baskerville', serif !important;
    font-size: 1.1rem !important;
    color: var(--bleu-marine) !important;
    font-weight: 700 !important;
    letter-spacing: 0.04em !important;
    text-transform: uppercase;
}
h3, h4, h5 {
    font-family: 'DM Sans', sans-serif !important;
    color: var(--bleu-moyen) !important;
    font-weight: 600 !important;
}

/* ---- Sidebar ---- */
[data-testid="stSidebar"] {
    background: var(--bleu-nuit) !important;
    border-right: 3px solid var(--or-medical) !important;
}
[data-testid="stSidebar"] * {
    color: var(--blanc-casse) !important;
}
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 {
    color: var(--or-clair) !important;
    font-family: 'DM Sans', sans-serif !important;
    text-transform: uppercase !important;
    font-size: 0.72rem !important;
    letter-spacing: 0.12em !important;
    border-bottom: 1px solid rgba(201,168,76,0.3) !important;
    padding-bottom: 0.4rem !important;
    margin-top: 1.2rem !important;
}
[data-testid="stSidebar"] .stRadio label,
[data-testid="stSidebar"] .stCheckbox label {
    color: rgba(245,242,238,0.85) !important;
    font-size: 0.88rem !important;
}
/* ---- File uploader sidebar : fix lisibilité complet ---- */
[data-testid="stSidebar"] [data-testid="stFileUploader"],
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] {
    background: rgba(255,255,255,0.06) !important;
    border: 1px dashed rgba(201,168,76,0.55) !important;
    border-radius: 8px !important;
    padding: 0.5rem !important;
}
/* Forcer blanc crème sur TOUT le contenu textuel du dropzone */
[data-testid="stSidebar"] [data-testid="stFileUploader"] *,
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] *,
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzoneInstructions"] *,
[data-testid="stSidebar"] section[data-testid="stFileUploaderDropzone"] span,
[data-testid="stSidebar"] section[data-testid="stFileUploaderDropzone"] small,
[data-testid="stSidebar"] section[data-testid="stFileUploaderDropzone"] div,
[data-testid="stSidebar"] section[data-testid="stFileUploaderDropzone"] p {
    color: var(--blanc-casse) !important;
    background: transparent !important;
}
/* Bouton "Browse files" */
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] button,
[data-testid="stSidebar"] section[data-testid="stFileUploaderDropzone"] button {
    background: rgba(201,168,76,0.15) !important;
    border: 1px solid rgba(201,168,76,0.5) !important;
    color: var(--or-clair) !important;
    border-radius: 6px !important;
}
[data-testid="stSidebar"] .uploadedFileName {
    color: var(--or-clair) !important;
}
[data-testid="stSidebar"] hr {
    border-color: rgba(201,168,76,0.25) !important;
}

/* ---- Inputs ---- */
input[type="number"],
input[type="text"],
input[type="password"],
.stTextInput input,
.stNumberInput input {
    background: white !important;
    border: 1.5px solid var(--gris-doux) !important;
    border-radius: 6px !important;
    color: var(--bleu-nuit) !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 0.92rem !important;
    transition: border-color 0.2s ease, box-shadow 0.2s ease !important;
}
input[type="number"]:focus,
input[type="text"]:focus,
.stTextInput input:focus,
.stNumberInput input:focus {
    border-color: var(--bleu-moyen) !important;
    box-shadow: 0 0 0 3px rgba(45,106,159,0.12) !important;
    outline: none !important;
}
.stNumberInput label,
.stTextInput label,
.stTextArea label,
.stRadio > label,
.stSelectbox label {
    font-size: 0.78rem !important;
    font-weight: 600 !important;
    color: var(--gris-texte) !important;
    text-transform: uppercase !important;
    letter-spacing: 0.06em !important;
}
.stTextArea textarea {
    background: white !important;
    border: 1.5px solid var(--gris-doux) !important;
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 0.9rem !important;
    color: var(--bleu-nuit) !important;
}
.stTextArea textarea:focus {
    border-color: var(--bleu-moyen) !important;
    box-shadow: 0 0 0 3px rgba(45,106,159,0.12) !important;
}

/* ---- Boutons ---- */
button[data-testid="baseButton-primary"],
.stButton > button[kind="primary"] {
    background: linear-gradient(135deg, var(--bleu-marine) 0%, var(--bleu-moyen) 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.95rem !important;
    letter-spacing: 0.04em !important;
    padding: 0.6rem 1.8rem !important;
    box-shadow: 0 4px 15px rgba(27,58,92,0.3) !important;
    transition: all 0.2s ease !important;
}
button[data-testid="baseButton-primary"]:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 20px rgba(27,58,92,0.4) !important;
}
.stButton > button[kind="secondary"] {
    background: transparent !important;
    color: var(--rouge-alerte) !important;
    border: 1.5px solid var(--rouge-alerte) !important;
    border-radius: 8px !important;
    font-weight: 500 !important;
    font-size: 0.88rem !important;
    transition: all 0.2s ease !important;
}
.stButton > button[kind="secondary"]:hover {
    background: var(--rouge-alerte) !important;
    color: white !important;
}

/* ---- Alertes ---- */
[data-testid="stAlert"][data-baseweb="notification"] {
    border-radius: 0 8px 8px 0 !important;
}
.stSuccess {
    background: rgba(46,125,90,0.1) !important;
    border-left: 4px solid var(--vert-ok) !important;
    border-radius: 0 8px 8px 0 !important;
}
.stWarning {
    background: rgba(201,168,76,0.12) !important;
    border-left: 4px solid var(--or-medical) !important;
    border-radius: 0 8px 8px 0 !important;
}
.stError {
    background: rgba(185,64,64,0.1) !important;
    border-left: 4px solid var(--rouge-alerte) !important;
    border-radius: 0 8px 8px 0 !important;
}
.stInfo {
    background: rgba(45,106,159,0.08) !important;
    border-left: 4px solid var(--bleu-moyen) !important;
    border-radius: 0 8px 8px 0 !important;
}

/* ---- Séparateurs ---- */
hr {
    border: none !important;
    border-top: 1px solid var(--gris-doux) !important;
    margin: 1.5rem 0 !important;
}

/* ---- Captions ---- */
.stCaption, small {
    color: var(--gris-texte) !important;
    font-size: 0.78rem !important;
    font-style: italic;
}

/* ---- Tabs ---- */
.stTabs [data-baseweb="tab-list"] {
    background: var(--gris-doux) !important;
    border-radius: 10px !important;
    padding: 4px !important;
    gap: 4px !important;
}
.stTabs [data-baseweb="tab"] {
    background: transparent !important;
    border-radius: 7px !important;
    color: var(--gris-texte) !important;
    font-weight: 500 !important;
    font-size: 0.9rem !important;
}
.stTabs [aria-selected="true"] {
    background: white !important;
    color: var(--bleu-marine) !important;
    font-weight: 600 !important;
    box-shadow: 0 1px 6px rgba(0,0,0,0.1) !important;
}

/* ---- Metric cards ---- */
[data-testid="metric-container"] {
    background: white !important;
    border-radius: 10px !important;
    padding: 1rem !important;
    border: 1px solid var(--gris-doux) !important;
    box-shadow: 0 2px 8px rgba(0,0,0,0.05) !important;
}

/* ---- Scrollbar ---- */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: var(--fond-app); }
::-webkit-scrollbar-thumb { background: var(--bleu-moyen); border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: var(--bleu-marine); }

/* ---- Responsive ---- */
@media (max-width: 768px) {
    h1 { font-size: 1.3rem !important; }
    h2 { font-size: 0.95rem !important; }
}
</style>
""", unsafe_allow_html=True)

# ==========================================
# 3. EN-TÊTE PROFESSIONNEL
# ==========================================
st.markdown("""
<div style="
    display: flex;
    align-items: center;
    gap: 1.2rem;
    padding: 0.5rem 0 1rem 0;
    margin-bottom: 0.5rem;
">
    <div style="font-size:2.6rem; line-height:1; filter: drop-shadow(0 2px 4px rgba(0,0,0,0.15));">🧠</div>
    <div>
        <div style="
            font-family: 'Libre Baskerville', serif;
            font-size: 1.7rem;
            font-weight: 700;
            color: #1B3A5C;
            letter-spacing: -0.02em;
            line-height: 1.15;
        ">Assistant WISC-V</div>
        <div style="
            font-size: 0.75rem;
            font-weight: 600;
            color: #C9A84C;
            text-transform: uppercase;
            letter-spacing: 0.12em;
            margin-top: 3px;
        ">Outil d'aide à l'analyse psychométrique</div>
    </div>
</div>
""", unsafe_allow_html=True)

st.warning("⚠️ **AVERTISSEMENT :** Outil d'aide à la rédaction. L'analyse clinique reste la responsabilité du psychologue.")

# ==========================================
# 4. SYSTÈME D'AUTHENTIFICATION
# ==========================================
def afficher_login():
    """Affiche la page de connexion stylée."""
    st.markdown("""
    <div style="
        max-width: 420px;
        margin: 4rem auto;
        background: white;
        border-radius: 16px;
        padding: 2.5rem 2rem;
        border: 1px solid #E8E4DF;
        border-top: 4px solid #C9A84C;
        box-shadow: 0 8px 40px rgba(0,0,0,0.10);
        text-align: center;
    ">
        <div style="font-size:2.8rem; margin-bottom:0.5rem;">🧠</div>
        <div style="
            font-family: 'Libre Baskerville', serif;
            font-size: 1.4rem;
            font-weight: 700;
            color: #1B3A5C;
            margin-bottom: 0.2rem;
        ">Assistant WISC-V</div>
        <div style="
            font-size: 0.72rem;
            color: #C9A84C;
            text-transform: uppercase;
            letter-spacing: 0.12em;
            font-weight: 600;
            margin-bottom: 2rem;
        ">Accès Professionnel Sécurisé</div>
    </div>
    """, unsafe_allow_html=True)

    col_vide1, col_form, col_vide2 = st.columns([1, 2, 1])
    with col_form:
        identifiant = st.text_input("👤 Identifiant", placeholder="Votre identifiant")
        mot_de_passe = st.text_input("🔑 Mot de passe", type="password", placeholder="Votre mot de passe")

        if st.button("Connexion", type="primary", use_container_width=True):
            if identifiant in UTILISATEURS:
                user = UTILISATEURS[identifiant]
                if user["hash"] == hash_mdp(mot_de_passe):
                    if date.today() > DATE_EXPIRATION:
                        st.error("⛔ Accès expiré. Contactez l'administrateur.")
                    else:
                        st.session_state.authenticated = True
                        st.session_state.user_id = identifiant
                        st.session_state.user_nom = user["nom"]
                        st.session_state.user_role = user["role"]
                        st.rerun()
                else:
                    st.error("❌ Mot de passe incorrect.")
            else:
                st.error("❌ Identifiant introuvable.")

        st.markdown("""
        <div style="
            text-align:center;
            font-size:0.75rem;
            color:#9CA3AF;
            margin-top:1.5rem;
        ">
            Outil réservé aux psychologues habilités.<br>
            Toute utilisation est tracée et soumise au secret professionnel.
        </div>
        """, unsafe_allow_html=True)

# Vérification de l'authentification
# (Désactivé pour l'instant — décommenter quand multi-utilisateurs nécessaire)
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = True
    st.session_state.user_nom = "Utilisateur"
    st.session_state.user_role = "admin"

# if not st.session_state.authenticated:
#     afficher_login()
#     st.stop()

try:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
except:
    st.error("⛔ Clé API manquante dans les secrets Streamlit.")
    st.stop()

# ==========================================
# 5. GESTION DU RESET
# ==========================================
if 'uploader_key' not in st.session_state:
    st.session_state.uploader_key = 0
if 'reset_confirm' not in st.session_state:
    st.session_state.reset_confirm = False

ALL_SCORE_KEYS = [
    'sim', 'voc', 'info', 'comp', 'cub', 'puz', 'mat', 'bal', 'arit',
    'memc', 'memi', 'seq', 'cod', 'sym', 'bar',
    'qit', 'perc_qit', 'qit_bas', 'qit_haut',
    'icv', 'perc_icv', 'icv_bas', 'icv_haut',
    'ivs', 'perc_ivs', 'ivs_bas', 'ivs_haut',
    'irf', 'perc_irf', 'irf_bas', 'irf_haut',
    'imt', 'perc_imt', 'imt_bas', 'imt_haut',
    'ivt', 'perc_ivt', 'ivt_bas', 'ivt_haut',
    'iag', 'iag_bas', 'iag_haut',
    'icc', 'icc_bas', 'icc_haut',
    'inv', 'inv_bas', 'inv_haut',
    'import_status'
]

def reset_all():
    for key in ALL_SCORE_KEYS:
        if key in st.session_state:
            del st.session_state[key]
    st.session_state['jn'] = 1; st.session_state['mn'] = 1; st.session_state['an'] = 2015
    st.session_state['jb'] = date.today().day
    st.session_state['mb'] = date.today().month
    st.session_state['ab'] = date.today().year
    st.session_state.uploader_key += 1
    st.session_state.reset_confirm = False
    st.rerun()

# ==========================================
# 6. INITIALISATION VARIABLES
# ==========================================
vars_scores = [
    'sim', 'voc', 'info', 'comp', 'cub', 'puz', 'mat', 'bal', 'arit',
    'memc', 'memi', 'seq', 'cod', 'sym', 'bar',
    'qit', 'perc_qit', 'qit_bas', 'qit_haut',
    'icv', 'perc_icv', 'icv_bas', 'icv_haut',
    'ivs', 'perc_ivs', 'ivs_bas', 'ivs_haut',
    'irf', 'perc_irf', 'irf_bas', 'irf_haut',
    'imt', 'perc_imt', 'imt_bas', 'imt_haut',
    'ivt', 'perc_ivt', 'ivt_bas', 'ivt_haut',
    'iag', 'iag_bas', 'iag_haut',
    'icc', 'icc_bas', 'icc_haut',
    'inv', 'inv_bas', 'inv_haut'
]
for var in vars_scores:
    if var not in st.session_state:
        st.session_state[var] = 0

if 'jn' not in st.session_state: st.session_state['jn'] = 1
if 'mn' not in st.session_state: st.session_state['mn'] = 1
if 'an' not in st.session_state: st.session_state['an'] = 2015
if 'jb' not in st.session_state: st.session_state['jb'] = date.today().day
if 'mb' not in st.session_state: st.session_state['mb'] = date.today().month
if 'ab' not in st.session_state: st.session_state['ab'] = date.today().year

# ==========================================
# 7. FONCTIONS
# ==========================================
def calculer_age(d_naiss, d_bilan):
    try:
        if d_bilan < d_naiss: return 0, 0
        ans = d_bilan.year - d_naiss.year
        mois = d_bilan.month - d_naiss.month
        if d_bilan.day < d_naiss.day: mois -= 1
        if mois < 0: ans -= 1; mois += 12
        return ans, mois
    except: return 0, 0

def check_homogeneite_indice(val1, val2, nom_indice):
    if val1 == 0 or val2 == 0: return None, ""
    ecart = abs(val1 - val2)
    if ecart >= 4: return False, f"⚠️ {nom_indice} Hétérogène (Écart {ecart})"
    else: return True, f"✅ {nom_indice} Homogène"

def plot_radar_chart(indices_dict):
    labels = list(indices_dict.keys())
    values = list(indices_dict.values())
    if sum(values) == 0: return None
    values += values[:1]
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]
    fig, ax = plt.subplots(figsize=(4, 4), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor('#F0EDE8')
    ax.set_facecolor('#F0EDE8')
    ax.fill(angles, values, color='#2D6A9F', alpha=0.2)
    ax.plot(angles, values, color='#1B3A5C', linewidth=2.5, label='Enfant')
    ax.plot(np.linspace(0, 2*np.pi, 100), [100]*100, color='#C9A84C',
            linestyle='--', linewidth=1.5, label='Norme (100)')
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=10, color='#1B3A5C', fontweight='600')
    ax.set_ylim(40, 160)
    ax.grid(color='#E8E4DF', linewidth=0.8)
    ax.spines['polar'].set_color('#E8E4DF')
    ax.legend(loc='upper right', bbox_to_anchor=(1.35, 1.15), fontsize='small')
    return fig

def read_file(file_obj, filename):
    text = ""
    try:
        if filename.lower().endswith('.pdf'):
            pdf_reader = PdfReader(file_obj)
            for page in pdf_reader.pages:
                t = page.extract_text()
                text += t + "\n" if t else ""
        else:
            stringio = StringIO(file_obj.getvalue().decode("utf-8"))
            text = stringio.read()
    except: pass
    return text

def extract_qglobal_data(text_content):
    try:
        model = genai.GenerativeModel('gemini-2.5-flash')
        prompt = f"""
        Extrais les données WISC-V du texte ci-dessous en JSON.
        IMPORTANT: Si une valeur est manquante, mets 0 (ou "" pour les dates).

        Variables Scores (Note Standard 1-19): sim, voc, info, comp, cub, puz, mat, bal, arit, memc, memi, seq, cod, sym, bar
        Indices (Note Composite): qit, icv, ivs, irf, imt, ivt, iag, icc, inv
        Percentiles: perc_qit, perc_icv, perc_ivs, perc_irf, perc_imt, perc_ivt
        IC 95% (Bas/Haut): qit_bas, qit_haut, icv_bas, icv_haut... (etc pour tous les indices).

        DATES (Format JJ/MM/AAAA) : date_naissance, date_passation

        TEXTE: {text_content[:9000]}
        Renvoie UNIQUEMENT un JSON valide.
        """
        response = model.generate_content(prompt)
        json_str = response.text.strip()
        if "```json" in json_str:
            json_str = json_str.split("```json")[1].split("```")[0]
        return json.loads(json_str)
    except Exception as e:
        st.error(f"Erreur extraction : {e}")
        return None

def create_docx(text_content, prenom, age_str):
    doc = Document()
    doc.add_heading(f'Compte Rendu WISC-V : {prenom}', 0)
    doc.add_paragraph(f"Âge au bilan : {age_str}")
    p = doc.add_paragraph()
    runner = p.add_run("AVERTISSEMENT : Document de travail. Analyse sous responsabilité du psychologue.")
    runner.bold = True; runner.italic = True
    doc.add_paragraph(text_content)
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def create_pdf(text_content, prenom, sexe, age_str, date_bilan_str):
    """Génère un PDF clinique professionnel avec mise en page soignée."""
    buf = io.BytesIO()

    # Couleurs
    BLEU_MARINE  = colors.HexColor('#1B3A5C')
    BLEU_MOYEN   = colors.HexColor('#2D6A9F')
    OR_MEDICAL   = colors.HexColor('#C9A84C')
    GRIS_DOUX    = colors.HexColor('#E8E4DF')
    GRIS_TEXTE   = colors.HexColor('#6B7280')
    FOND_CLAIR   = colors.HexColor('#F5F2EE')

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        title=f"Bilan WISC-V - {prenom}",
        author="Assistant WISC-V"
    )

    # Styles
    styles = getSampleStyleSheet()

    s_titre = ParagraphStyle('titre',
        fontSize=18, textColor=BLEU_MARINE,
        fontName='Helvetica-Bold', alignment=TA_LEFT,
        spaceAfter=4)

    s_sous_titre = ParagraphStyle('sous_titre',
        fontSize=9, textColor=OR_MEDICAL,
        fontName='Helvetica-Bold', alignment=TA_LEFT,
        spaceAfter=16, letterSpacing=1.5)

    s_section = ParagraphStyle('section',
        fontSize=10, textColor=BLEU_MARINE,
        fontName='Helvetica-Bold', alignment=TA_LEFT,
        spaceBefore=14, spaceAfter=6,
        borderPad=4)

    s_corps = ParagraphStyle('corps',
        fontSize=9.5, textColor=colors.HexColor('#1a2a3a'),
        fontName='Helvetica', alignment=TA_JUSTIFY,
        spaceAfter=8, leading=15)

    s_avert = ParagraphStyle('avert',
        fontSize=8, textColor=GRIS_TEXTE,
        fontName='Helvetica-Oblique', alignment=TA_CENTER,
        spaceAfter=6)

    s_pied = ParagraphStyle('pied',
        fontSize=7.5, textColor=GRIS_TEXTE,
        fontName='Helvetica', alignment=TA_CENTER)

    story = []

    # --- En-tête ---
    story.append(Paragraph("Assistant WISC-V", s_titre))
    story.append(Paragraph("COMPTE RENDU PSYCHOMÉTRIQUE", s_sous_titre))
    story.append(HRFlowable(width="100%", thickness=2, color=OR_MEDICAL, spaceAfter=12))

    # --- Tableau identité ---
    data_id = [
        ["Prénom", prenom,          "Sexe",       sexe],
        ["Âge au bilan", age_str,   "Date bilan", date_bilan_str],
    ]
    t = Table(data_id, colWidths=[3.5*cm, 6*cm, 3.5*cm, 6*cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), FOND_CLAIR),
        ('BACKGROUND', (2,0), (2,-1), FOND_CLAIR),
        ('TEXTCOLOR',  (0,0), (0,-1), BLEU_MARINE),
        ('TEXTCOLOR',  (2,0), (2,-1), BLEU_MARINE),
        ('FONTNAME',   (0,0), (0,-1), 'Helvetica-Bold'),
        ('FONTNAME',   (2,0), (2,-1), 'Helvetica-Bold'),
        ('FONTSIZE',   (0,0), (-1,-1), 9),
        ('GRID',       (0,0), (-1,-1), 0.5, GRIS_DOUX),
        ('PADDING',    (0,0), (-1,-1), 6),
        ('ROWBACKGROUNDS', (0,0), (-1,-1), [colors.white, colors.white]),
    ]))
    story.append(t)
    story.append(Spacer(1, 14))

    # --- Avertissement ---
    story.append(HRFlowable(width="100%", thickness=0.5, color=GRIS_DOUX))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        "⚠ Document de travail confidentiel — L'analyse clinique et les conclusions restent "
        "sous la responsabilité exclusive du psychologue.",
        s_avert))
    story.append(HRFlowable(width="100%", thickness=0.5, color=GRIS_DOUX))
    story.append(Spacer(1, 14))

    # --- Contenu de l'analyse ---
    # Découper le texte en sections selon les titres markdown
    lignes = text_content.split('\n')
    for ligne in lignes:
        ligne = ligne.strip()
        if not ligne:
            story.append(Spacer(1, 4))
        elif ligne.startswith('## ') or ligne.startswith('# '):
            titre_propre = ligne.lstrip('#').strip()
            story.append(HRFlowable(width="100%", thickness=0.5, color=GRIS_DOUX, spaceBefore=8))
            story.append(Paragraph(titre_propre.upper(), s_section))
        elif ligne.startswith('### '):
            titre_propre = ligne.lstrip('#').strip()
            story.append(Paragraph(f"<b>{titre_propre}</b>", s_corps))
        elif ligne.startswith('**') and ligne.endswith('**'):
            story.append(Paragraph(f"<b>{ligne.strip('*')}</b>", s_corps))
        elif ligne.startswith('- ') or ligne.startswith('* '):
            story.append(Paragraph(f"• {ligne[2:]}", s_corps))
        else:
            # Nettoyer les astérisques markdown restants
            ligne_clean = ligne.replace('**', '<b>', 1).replace('**', '</b>', 1)
            story.append(Paragraph(ligne_clean, s_corps))

    # --- Pied de page ---
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=1, color=OR_MEDICAL))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        f"Document généré le {date.today().strftime('%d/%m/%Y')} · Assistant WISC-V · "
        "Confidentiel — Secret professionnel",
        s_pied))

    doc.build(story)
    buf.seek(0)
    return buf

# ==========================================
# 8. SIDEBAR
# ==========================================
knowledge_base = ""
with st.sidebar:
    # --- Utilisateur connecté ---
    st.markdown(f"""
    <div style="
        text-align:center;
        padding: 1rem 0 0.8rem 0;
        border-bottom: 1px solid rgba(201,168,76,0.3);
        margin-bottom: 0.5rem;
    ">
        <div style="font-size:1.8rem;">🧠</div>
        <div style="
            font-size:0.65rem;
            color: #C9A84C;
            text-transform:uppercase;
            letter-spacing:0.12em;
            font-weight:600;
            margin-top:4px;
        ">WISC-V · Assistant Pro</div>
        <div style="
            font-size:0.78rem;
            color: rgba(245,242,238,0.7);
            margin-top:8px;
        ">👤 {st.session_state.get('user_nom', '')}
        {'🔧' if st.session_state.get('user_role') == 'admin' else ''}
        </div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("🚪 Déconnexion", use_container_width=True):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

    st.header("⚙️ Configuration")
    style_redac = st.radio(
        "Destinataire / Style",
        ["Expert / MDPH (Technique & Clinique)", "Parents / École (Pédagogique)"],
        index=0
    )

    st.divider()
    st.header("📥 Import Q-GLOBAL")
    uploaded_qglobal = st.file_uploader(
        "Rapport PDF",
        type=['pdf', 'txt'],
        key=f"uploader_{st.session_state.uploader_key}"
    )

    if 'import_status' in st.session_state:
        status = st.session_state['import_status']
        if status['success']:
            st.success(status['msg'])
            if status['missing']:
                st.warning(f"⚠️ Manquant (mis à 0) :\n" + ", ".join(status['missing']))
        else:
            st.error(status['msg'])

    if uploaded_qglobal and st.button("🚀 Extraire Données", type="primary"):
        with st.spinner("Analyse IA en cours..."):
            raw = read_file(uploaded_qglobal, uploaded_qglobal.name)
            data_ex = extract_qglobal_data(raw)
            missing = []; count = 0
            if data_ex:
                for k, v in data_ex.items():
                    if k == 'date_naissance' and v:
                        try:
                            d = v.split('/')
                            st.session_state['jn'] = int(d[0])
                            st.session_state['mn'] = int(d[1])
                            st.session_state['an'] = int(d[2])
                            count += 1
                        except: pass
                    elif k == 'date_passation' and v:
                        try:
                            d = v.split('/')
                            st.session_state['jb'] = int(d[0])
                            st.session_state['mb'] = int(d[1])
                            st.session_state['ab'] = int(d[2])
                            count += 1
                        except: pass
                    elif k in st.session_state:
                        try:
                            if v is None or v == "":
                                val = 0; missing.append(k)
                            else:
                                val = float(v)
                            if val == 0 and k not in missing:
                                missing.append(k)
                            if 'perc' in k:
                                st.session_state[k] = val
                            else:
                                st.session_state[k] = int(val)
                            count += 1
                        except:
                            st.session_state[k] = 0
                            missing.append(k)
                st.session_state['import_status'] = {
                    'success': True,
                    'msg': f"✅ {count} champs importés.",
                    'missing': missing
                }
                st.rerun()
            else:
                st.session_state['import_status'] = {
                    'success': False,
                    'msg': "Échec extraction IA.",
                    'missing': []
                }

    st.divider()
    st.header("📚 Bibliothèque")
    local_files = [
        f for f in os.listdir('.')
        if f.lower().endswith(('.pdf', '.txt'))
        and f not in ["requirements.txt", "app.py"]
    ]
    if local_files:
        for f in local_files:
            if st.checkbox(f"📄 {f}", value=True, key=f):
                c = read_file(f, f)
                knowledge_base += f"\n--- SOURCE PRIORITAIRE: {f} ---\n{c}\n"
        st.caption(f"Contexte : {len(knowledge_base)} chars")
        with st.expander("👀 Vérifier le contenu lu par l'IA"):
            st.text(knowledge_base[:3000] + "...")
    else:
        st.warning("Pas de PDF trouvés.")

    st.divider()
    if not st.session_state.reset_confirm:
        if st.button("🗑️ Nouvelle Analyse (Reset)", type="secondary"):
            st.session_state.reset_confirm = True
            st.rerun()
    else:
        st.warning("⚠️ Toutes les données seront effacées. Confirmer ?")
        col_oui, col_non = st.columns(2)
        with col_oui:
            if st.button("✅ Oui, effacer", type="primary"):
                reset_all()
        with col_non:
            if st.button("❌ Annuler"):
                st.session_state.reset_confirm = False
                st.rerun()

# ==========================================
# 9. INTERFACE PRINCIPALE
# ==========================================

# --- Section 1 : Identité ---
st.markdown("""
<div style="
    background: white;
    border-radius: 12px;
    padding: 1.2rem 1.5rem 0.5rem 1.5rem;
    border: 1px solid #E8E4DF;
    box-shadow: 0 2px 12px rgba(0,0,0,0.05);
    margin-bottom: 1rem;
">
<h2 style="margin-top:0 !important;">1. Identité</h2>
""", unsafe_allow_html=True)

c1, c2, c3 = st.columns(3)
with c1:
    prenom = st.text_input("Prénom", placeholder="Ex : Lucas")
    sexe = st.radio("Sexe", ["Garçon", "Fille"], horizontal=True)
    lateralite = st.radio("Latéralité", ["Droitier", "Gaucher"], horizontal=True)
with c2:
    st.markdown("**Date de naissance**")
    cj, cm, ca = st.columns([1, 1, 1.5])
    with cj: jn = st.number_input("J", 1, 31, key="jn")
    with cm: mn = st.number_input("M", 1, 12, key="mn")
    with ca: an = st.number_input("A", 2000, 2030, key="an")
    try: dn = date(an, mn, jn)
    except: dn = date.today()
with c3:
    st.markdown("**Date du bilan**")
    cj, cm, ca = st.columns([1, 1, 1.5])
    with cj: jb = st.number_input("J", 1, 31, key="jb")
    with cm: mb = st.number_input("M", 1, 12, key="mb")
    with ca: ab = st.number_input("A", 2020, 2030, key="ab")
    try: dt = date(ab, mb, jb)
    except: dt = date.today()
    ans, mois = calculer_age(dn, dt)
    st.success(f"**{ans} ans {mois} mois**")

st.markdown("</div>", unsafe_allow_html=True)

# --- Section 2 : Clinique ---
st.markdown("""
<div style="
    background: white;
    border-radius: 12px;
    padding: 1.2rem 1.5rem 0.5rem 1.5rem;
    border: 1px solid #E8E4DF;
    box-shadow: 0 2px 12px rgba(0,0,0,0.05);
    margin-bottom: 1rem;
">
<h2 style="margin-top:0 !important;">2. Observations Cliniques</h2>
""", unsafe_allow_html=True)

c1, c2, c3 = st.columns(3)
obs = []
with c1:
    st.markdown("**Attitude**")
    if st.checkbox("Anxiété"): obs.append("Anxiété perf.")
    if st.checkbox("Opposition"): obs.append("Opposition")
    if st.checkbox("Agitation"): obs.append("Agitation")
    if st.checkbox("Impulsivité"): obs.append("Impulsivité")
with c2:
    st.markdown("**Cognition**")
    if st.checkbox("Fatigabilité"): obs.append("Fatigabilité")
    if st.checkbox("Inattention"): obs.append("Inattention")
    if st.checkbox("Besoin relance"): obs.append("Besoin relance")
    if st.checkbox("Verbal +++"): obs.append("Logorrhée")
    if st.checkbox("Verbal ---"): obs.append("Mutisme/Pauvreté")
with c3:
    st.markdown("**Graphisme**")
    if st.checkbox("Crispation"): obs.append("Crispation")
    if st.checkbox("Lenteur graph."): obs.append("Lenteur graph.")
    st.markdown("---")
    st.markdown("🗣️ **Langue / Créole**")
    creole = st.radio(
        "Usage créole",
        ["-- (Non/Peu)", "+- (Moyen)", "++ (Dominant)"],
        index=0,
        label_visibility="collapsed"
    )

st.markdown("**🎯 Motif de consultation**")
motifs_col1, motifs_col2 = st.columns(2)
motifs = []
with motifs_col1:
    if st.checkbox("Difficultés scolaires"): motifs.append("Difficultés scolaires")
    if st.checkbox("Suspicion TDAH"): motifs.append("Suspicion TDAH")
    if st.checkbox("Suspicion TSA"): motifs.append("Suspicion TSA")
    if st.checkbox("Suspicion HPI / Douance"): motifs.append("Suspicion HPI/Douance")
with motifs_col2:
    if st.checkbox("Orientation MDPH / RQTH"): motifs.append("Orientation MDPH/RQTH")
    if st.checkbox("Bilan de rééducation"): motifs.append("Bilan de rééducation")
    if st.checkbox("Suivi psy / thérapeutique"): motifs.append("Suivi psy/thérapeutique")
    if st.checkbox("Autre motif"): motifs.append("Autre (voir anamnèse)")
st.markdown("---")
obs_libre = st.text_area("Observations libres", height=70)

st.markdown("</div>", unsafe_allow_html=True)

# --- Section 3 : Psychométrie ---
st.markdown("""
<div style="
    background: white;
    border-radius: 12px;
    padding: 1.2rem 1.5rem 1rem 1.5rem;
    border: 1px solid #E8E4DF;
    box-shadow: 0 2px 12px rgba(0,0,0,0.05);
    margin-bottom: 1rem;
">
<h2 style="margin-top:0 !important;">3. Psychométrie</h2>
""", unsafe_allow_html=True)

st.subheader("A. Subtests (Notes Standard)")

def couleur_note(val):
    """Retourne une couleur selon la note standard."""
    if val == 0: return ""
    if val <= 6:   return "🔴"
    if val <= 9:   return "🟠"
    if val <= 12:  return "🟡"
    if val <= 15:  return "🟢"
    return "🔵"

c1, c2, c3, c4 = st.columns(4)
with c1: sim  = st.number_input(f"SIM {couleur_note(st.session_state.get('sim',0))}", 0, 19, key="sim")
with c2: voc  = st.number_input(f"VOC {couleur_note(st.session_state.get('voc',0))}", 0, 19, key="voc")
with c3: info = st.number_input(f"INF {couleur_note(st.session_state.get('info',0))}", 0, 19, key="info")
with c4: comp = st.number_input(f"COM {couleur_note(st.session_state.get('comp',0))}", 0, 19, key="comp")

c1, c2 = st.columns(2)
with c1: cub = st.number_input(f"CUB {couleur_note(st.session_state.get('cub',0))}", 0, 19, key="cub")
with c2: puz = st.number_input(f"PUZ {couleur_note(st.session_state.get('puz',0))}", 0, 19, key="puz")

c1, c2, c3 = st.columns(3)
with c1: mat  = st.number_input(f"MAT {couleur_note(st.session_state.get('mat',0))}", 0, 19, key="mat")
with c2: bal  = st.number_input(f"BAL {couleur_note(st.session_state.get('bal',0))}", 0, 19, key="bal")
with c3: arit = st.number_input(f"ARI {couleur_note(st.session_state.get('arit',0))}", 0, 19, key="arit")

c1, c2, c3 = st.columns(3)
with c1: memc = st.number_input(f"MCH {couleur_note(st.session_state.get('memc',0))}", 0, 19, key="memc")
with c2: memi = st.number_input(f"MIM {couleur_note(st.session_state.get('memi',0))}", 0, 19, key="memi")
with c3: seq  = st.number_input(f"SLC {couleur_note(st.session_state.get('seq',0))}", 0, 19, key="seq")

c1, c2, c3 = st.columns(3)
with c1: cod = st.number_input(f"COD {couleur_note(st.session_state.get('cod',0))}", 0, 19, key="cod")
with c2: sym = st.number_input(f"SYM {couleur_note(st.session_state.get('sym',0))}", 0, 19, key="sym")
with c3: bar = st.number_input(f"BAR {couleur_note(st.session_state.get('bar',0))}", 0, 19, key="bar")

st.caption("🔴 ≤6 Très faible · 🟠 7-9 Faible · 🟡 10-12 Moyen · 🟢 13-15 Supérieur · 🔵 ≥16 Très supérieur")

st.markdown("---")
st.subheader("B. Indices (Note / Percentile / Intervalle de confiance)")

# Homogénéité
vicv, ticv = check_homogeneite_indice(sim, voc, "ICV")
vivs, tivs = check_homogeneite_indice(cub, puz, "IVS")
virf, tirf = check_homogeneite_indice(mat, bal, "IRF")
vimt, timt = check_homogeneite_indice(memc, memi, "IMT")
vivt, tivt = check_homogeneite_indice(sym, cod, "IVT")
nb_inv = sum([1 for x in [vicv, vivs, virf, vimt, vivt] if x is False])

# --- Indices principaux (en premier) ---
c1, c2, c3, c4, c5 = st.columns(5)
with c1:
    st.markdown("**ICV**")
    icv = st.number_input("Note ICV", 0, 160, key="icv", label_visibility="collapsed")
    perc_icv = st.number_input("P_ICV", 0, 100, key="perc_icv", label_visibility="collapsed")
    icv_bas = st.number_input("IC Bas ICV", 0, 160, key="icv_bas")
    icv_haut = st.number_input("IC Haut ICV", 0, 160, key="icv_haut")
    if ticv: st.caption(ticv)
with c2:
    st.markdown("**IVS**")
    ivs = st.number_input("Note IVS", 0, 160, key="ivs", label_visibility="collapsed")
    perc_ivs = st.number_input("P_IVS", 0, 100, key="perc_ivs", label_visibility="collapsed")
    ivs_bas = st.number_input("IC Bas IVS", 0, 160, key="ivs_bas")
    ivs_haut = st.number_input("IC Haut IVS", 0, 160, key="ivs_haut")
    if tivs: st.caption(tivs)
with c3:
    st.markdown("**IRF**")
    irf = st.number_input("Note IRF", 0, 160, key="irf", label_visibility="collapsed")
    perc_irf = st.number_input("P_IRF", 0, 100, key="perc_irf", label_visibility="collapsed")
    irf_bas = st.number_input("IC Bas IRF", 0, 160, key="irf_bas")
    irf_haut = st.number_input("IC Haut IRF", 0, 160, key="irf_haut")
    if tirf: st.caption(tirf)
with c4:
    st.markdown("**IMT**")
    imt = st.number_input("Note IMT", 0, 160, key="imt", label_visibility="collapsed")
    perc_imt = st.number_input("P_IMT", 0, 100, key="perc_imt", label_visibility="collapsed")
    imt_bas = st.number_input("IC Bas IMT", 0, 160, key="imt_bas")
    imt_haut = st.number_input("IC Haut IMT", 0, 160, key="imt_haut")
    if timt: st.caption(timt)
with c5:
    st.markdown("**IVT**")
    ivt = st.number_input("Note IVT", 0, 160, key="ivt", label_visibility="collapsed")
    perc_ivt = st.number_input("P_IVT", 0, 100, key="perc_ivt", label_visibility="collapsed")
    ivt_bas = st.number_input("IC Bas IVT", 0, 160, key="ivt_bas")
    ivt_haut = st.number_input("IC Haut IVT", 0, 160, key="ivt_haut")
    if tivt: st.caption(tivt)

# --- Validité globale ---
st.markdown("---")
chk = [icv, ivs, irf, imt, ivt]
if all(i > 0 for i in chk):
    disp = max(chk) - min(chk)
    if disp >= 23:
        st.error(f"⚠️ **QIT NON INTERPRÉTABLE** — Dispersion entre indices : {disp} points")
        h_txt = f"NON INTERPRÉTABLE (Disp. {disp})"
    elif nb_inv >= 2:
        st.warning(f"🟠 **QIT FRAGILE** — {nb_inv} indice(s) hétérogène(s)")
        h_txt = f"FRAGILE ({nb_inv} ind. hétérogènes)"
    else:
        st.success("✅ **Profil homogène** — QIT interprétable")
        h_txt = "Valide et Homogène"
else:
    st.info("ℹ️ Renseignez les 5 indices pour évaluer la validité du QIT")
    h_txt = "N/A"

# --- QIT en dernier ---
st.markdown("**QIT (Quotient Intellectuel Total)**")
col_qit1, col_qit2, col_qit3, col_qit4 = st.columns(4)
with col_qit1: qit = st.number_input("Note QIT", 0, 160, key="qit")
with col_qit2: perc_qit = st.number_input("Percentile", 0, 100, key="perc_qit")
with col_qit3: qit_bas = st.number_input("IC Bas QIT", 0, 160, key="qit_bas")
with col_qit4: qit_haut = st.number_input("IC Haut QIT", 0, 160, key="qit_haut")

st.markdown("---")
st.subheader("C. Indices Complémentaires")

def safe_sum(values):
    if all(v > 0 for v in values): return sum(values)
    return "Incomplet"

s_iag = safe_sum([sim, voc, cub, mat, bal])
s_icc = safe_sum([memc, memi, sym, cod])
s_inv = safe_sum([cub, puz, mat, bal, memi, cod])

st.caption(f"🧮 **Aide calcul (Somme Notes Standard) :** IAG = **{s_iag}** | ICC = **{s_icc}** | INV = **{s_inv}**")

c1, c2, c3 = st.columns(3)
with c1:
    st.markdown("**IAG**")
    iag = st.number_input("IAG", 0, key="iag")
    iag_bas = st.number_input("IB_IAG", 0, key="iag_bas", label_visibility="collapsed")
    iag_haut = st.number_input("IH_IAG", 0, key="iag_haut", label_visibility="collapsed")
with c2:
    st.markdown("**ICC**")
    icc = st.number_input("ICC", 0, key="icc")
    icc_bas = st.number_input("IB_ICC", 0, key="icc_bas", label_visibility="collapsed")
    icc_haut = st.number_input("IH_ICC", 0, key="icc_haut", label_visibility="collapsed")
with c3:
    st.markdown("**INV**")
    inv = st.number_input("INV", 0, key="inv")
    inv_bas = st.number_input("IB_INV", 0, key="inv_bas", label_visibility="collapsed")
    inv_haut = st.number_input("IH_INV", 0, key="inv_haut", label_visibility="collapsed")

st.markdown("</div>", unsafe_allow_html=True)

# ==========================================
# 10. ANALYSE & GRAPHIQUE
# ==========================================
st.divider()
indices = {"ICV": icv, "IVS": ivs, "IRF": irf, "IMT": imt, "IVT": ivt}
valid_ind = {k: v for k, v in indices.items() if v > 0}

c1, c2 = st.columns([1, 1.5])
with c1:
    if len(valid_ind) >= 3:
        st.pyplot(plot_radar_chart(valid_ind))
with c2:
    if valid_ind:
        vals = list(valid_ind.values())
        moy = np.mean(vals)
        et = np.std(vals)
        st.info(f"Moyenne Perso : **{moy:.1f}** | Écart-Type : **{et:.1f}**")
        intra_txt = f"Moyenne Perso: {moy:.1f}, ET: {et:.1f}."
        for k, v in valid_ind.items():
            d = v - moy
            if d >= 10:
                st.write(f"🟢 **{k}** : Force relative (+{d:.1f})")
                intra_txt += f"- {k}: Force relative.\n"
            elif d <= -10:
                st.write(f"🔴 **{k}** : Faiblesse relative ({d:.1f})")
                intra_txt += f"- {k}: Faiblesse relative.\n"
    else:
        intra_txt = ""

# ==========================================
# 11. GÉNÉRATION IA
# ==========================================
st.divider()

# Point 2 : Avertissement si bibliothèque vide
if not knowledge_base:
    st.error("⛔ **Bibliothèque vide !** Aucun PDF de référence n'est chargé dans la sidebar. L'analyse sera moins précise sur le plan méthodologique.")

if st.button("✨ GÉNÉRER L'ANALYSE EXPERT", type="primary"):
    infos = f"{prenom}, {sexe}, {ans} ans. Date Bilan: {st.session_state.jb}/{st.session_state.mb}/{st.session_state.ab}. Latéralité: {lateralite}. Créole: {creole}."
    motif_txt = ", ".join(motifs) if motifs else "Non précisé"
    obs_txt = ", ".join(obs) + ". " + obs_libre

    if qit > 0:
        data = f"QIT: {qit} (Perc: {perc_qit}, IC: {qit_bas}-{qit_haut}). Validité: {h_txt}.\n"
    else:
        data = "QIT: Non calculé / Non administré.\n"

    data += "Indices Administrés: "
    indices_data = []
    if icv > 0: indices_data.append(f"ICV {icv} (Perc {perc_icv}, IC {icv_bas}-{icv_haut})")
    if ivs > 0: indices_data.append(f"IVS {ivs} (Perc {perc_ivs}, IC {ivs_bas}-{ivs_haut})")
    if irf > 0: indices_data.append(f"IRF {irf} (Perc {perc_irf}, IC {irf_bas}-{irf_haut})")
    if imt > 0: indices_data.append(f"IMT {imt} (Perc {perc_imt}, IC {imt_bas}-{imt_haut})")
    if ivt > 0: indices_data.append(f"IVT {ivt} (Perc {perc_ivt}, IC {ivt_bas}-{ivt_haut})")
    data += ", ".join(indices_data) + ".\n"

    data += "Indices Complémentaires: "
    compl_data = []
    if iag > 0: compl_data.append(f"IAG {iag} (IC {iag_bas}-{iag_haut})")
    if icc > 0: compl_data.append(f"ICC {icc} (IC {icc_bas}-{icc_haut})")
    if inv > 0: compl_data.append(f"INV {inv} (IC {inv_bas}-{inv_haut})")
    data += ", ".join(compl_data) + ".\n"

    data += "Subtests (Notes Standard): "
    subs_map = {
        "Sim": sim, "Voc": voc, "Info": info, "Comp": comp,
        "Cub": cub, "Puz": puz, "Mat": mat, "Bal": bal,
        "Arit": arit, "MemC": memc, "MemI": memi, "Seq": seq,
        "Cod": cod, "Sym": sym, "Bar": bar
    }
    valid_subs = [f"{k} {v}" for k, v in subs_map.items() if v > 0]
    data += ", ".join(valid_subs) + "."

    with st.spinner("🧠 Analyse approfondie en cours..."):
        try:
            model = genai.GenerativeModel('gemini-2.5-flash')
            prompt = f"""
            Rôle: Expert Psychologue WISC-V.

            DESTINATAIRE: {style_redac}.

            DONNÉES ENTRÉE:
            - Enfant: {infos}
            - Motif de consultation: {motif_txt}
            - Obs: {obs_txt}
            - Anamnèse: {ana}
            - Scores: {data}
            - Stats Intra: {intra_txt}

            <BIBLIOTHEQUE_REFERENCE>
            {knowledge_base}
            </BIBLIOTHEQUE_REFERENCE>

            CONSIGNE CRUCIALE DE HIERARCHIE :
            1. Pour la MÉTHODOLOGIE (calculs, validité, homogénéité), tu DOIS suivre scrupuleusement le contenu de <BIBLIOTHEQUE_REFERENCE> (notamment Grégoire, Ozenne).
            2. Pour le VOCABULAIRE DIAGNOSTIQUE en conclusion, utilise le DSM-5 / CIM-11.

            OBJECTIF QUALITATIF :
            Ne te contente pas de lister les scores. Tu dois EXPLIQUER et INTERPRÉTER.
            - Utilise des connecteurs logiques : "ce qui suggère que...", "probablement en raison de...", "ce résultat contraste avec...".
            - Formule des HYPOTHÈSES sur les mécanismes cognitifs sous-jacents.
            - Analyse les ÉCARTS : Si l'ICV est > IVS, qu'est-ce que ça implique concrètement ?
            - Tiens compte du motif de consultation pour orienter la conclusion et les recommandations.

            STRUCTURE DU COMPTE RENDU :

            1. VALIDITÉ DES INDICES GLOBAUX
               - Vérifier homogénéité QIT. Si invalide, passer à IAG/ICC/INV.

            2. ANALYSE INTER-INDIVIDUELLE (NORMATIVE) -> FOCUS INDICES UNIQUEMENT
               - Parle des INDICES (QIT, ICV, etc.) par rapport à la norme (100).
               - INTERDICTION de parler des subtests ici.
               *** SYNTHÈSE NORMATIVE & FONCTIONNELLE ***
               - Paragraphe de synthèse sur l'efficience globale et l'impact sur la vie quotidienne/scolaire.

            3. ANALYSE INTRA-INDIVIDUELLE (IPSATIVE) -> FOCUS SUBTESTS
               - Analyse les SUBTESTS par rapport à la moyenne de l'enfant ({moy if valid_ind else 'N/A'}).
               - Lier chaque résultat à une hypothèse cognitive/clinique.
               *** SYNTHÈSE CLINIQUE & PROCESSUELLE ***
               - Forces et faiblesses spécifiques + lien avec les symptômes observés.

            4. SYNTHÈSE DIAGNOSTIQUE & RECOMMANDATIONS
               - Croiser avec l'anamnèse et le motif de consultation.
               - Hypothèses (TDAH, TSA, etc.).
               - Conseils pratiques adaptés au motif.

            Rédige avec un ton professionnel, argumenté et clinique.
            """

            res = model.generate_content(prompt)
            analyse_texte = res.text

            # Stocker le résultat pour le résumé
            st.session_state['derniere_analyse'] = analyse_texte
            st.session_state['prenom_analyse'] = prenom
            st.session_state['age_analyse'] = f"{ans}a{mois}m"

            st.markdown("""
            <div style="
                background: white;
                border-radius: 12px;
                padding: 1.5rem 2rem;
                border: 1px solid #E8E4DF;
                border-left: 4px solid #1B3A5C;
                box-shadow: 0 4px 20px rgba(0,0,0,0.06);
                margin-top: 1rem;
            ">
            """, unsafe_allow_html=True)
            st.markdown(analyse_texte)
            st.markdown("</div>", unsafe_allow_html=True)

            # Export Word + PDF côte à côte
            col_word, col_pdf = st.columns(2)
            with col_word:
                f_word = create_docx(analyse_texte, prenom, f"{ans}a{mois}m")
                st.download_button(
                    "📄 Télécharger (.docx)",
                    f_word,
                    f"Bilan_WISC5_{prenom}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            with col_pdf:
                date_bilan_str = f"{st.session_state.jb:02d}/{st.session_state.mb:02d}/{st.session_state.ab}"
                f_pdf = create_pdf(analyse_texte, prenom, sexe, f"{ans} ans {mois} mois", date_bilan_str)
                st.download_button(
                    "📋 Télécharger (.pdf)",
                    f_pdf,
                    f"Bilan_WISC5_{prenom}.pdf",
                    "application/pdf",
                    use_container_width=True
                )

        except Exception as e:
            st.error(f"Erreur lors de la génération : {e}")

# ==========================================
# 12. RÉSUMÉ 10 LIGNES (Point 8)
# ==========================================
if 'derniere_analyse' in st.session_state:
    st.divider()
    st.markdown("### 📋 Résumé synthétique")
    st.caption("Utile pour courriers, transmissions MDPH, comptes rendus rapides.")
    if st.button("✍️ Générer un résumé en 10 lignes maximum", type="secondary"):
        with st.spinner("Rédaction du résumé..."):
            try:
                model = genai.GenerativeModel('gemini-2.5-flash')
                prompt_resume = f"""
                À partir de cette analyse WISC-V complète, rédige un résumé synthétique 
                destiné à un courrier professionnel (médecin, école, MDPH).
                
                Contraintes STRICTES :
                - 10 lignes MAXIMUM
                - Ton professionnel, phrases complètes
                - Inclure : efficience globale, points forts, points faibles, 1-2 recommandations clés
                - Ne pas utiliser de titres ni de bullet points, uniquement des paragraphes
                - Commencer par "À l'issue du bilan psychométrique de {st.session_state.get('prenom_analyse','cet enfant')}..."
                
                ANALYSE SOURCE :
                {st.session_state['derniere_analyse']}
                """
                res_resume = model.generate_content(prompt_resume)

                st.markdown("""
                <div style="
                    background: #F5F2EE;
                    border-radius: 10px;
                    padding: 1.2rem 1.5rem;
                    border: 1px solid #E8E4DF;
                    border-left: 4px solid #C9A84C;
                    margin-top: 0.5rem;
                ">
                """, unsafe_allow_html=True)
                st.markdown(res_resume.text)
                st.markdown("</div>", unsafe_allow_html=True)

                # Bouton copier (via text_area sélectionnable)
                st.text_area(
                    "📋 Sélectionner tout pour copier :",
                    res_resume.text,
                    height=200,
                    key="resume_copie"
                )

            except Exception as e:
                st.error(f"Erreur résumé : {e}")
